"""Extraccion de texto de un CV subido (PDF / DOCX / TXT).

Solo extrae TEXTO; la estructuracion a cv_master la hace `ai.profile_extractor`
con el LLM. PDFs con columnas/tablas/escaneados se extraen mal: por eso siempre
hay una pantalla de revision editable aguas abajo. Devuelve (texto, warnings).

Dependencias en el extra `onboarding`: pypdf, pdfplumber, python-docx.
"""

from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)

# Umbral por debajo del cual sospechamos PDF escaneado / extraccion pobre.
_MIN_CHARS = 200


def _extract_pdf(data: bytes) -> tuple[str, list[str]]:
    warnings: list[str] = []
    text = ""
    # 1) pypdf (rapido)
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    except Exception as exc:  # noqa: BLE001
        warnings.append(f"pypdf fallo: {exc}")

    # 2) pdfplumber (mejor con layout) si pypdf saco poco
    if len(text) < _MIN_CHARS:
        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(data)) as pdf:
                text = "\n".join((p.extract_text() or "") for p in pdf.pages).strip()
        except Exception as exc:  # noqa: BLE001
            warnings.append(f"pdfplumber fallo: {exc}")

    if len(text) < _MIN_CHARS:
        warnings.append(
            "Se extrajo muy poco texto del PDF (posible escaneado o con columnas). "
            "Revisa y completa el perfil manualmente."
        )
    return text, warnings


def _extract_docx(data: bytes) -> tuple[str, list[str]]:
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts).strip(), []
    except Exception as exc:  # noqa: BLE001
        return "", [f"docx fallo: {exc}"]


def extract_text(filename: str, data: bytes) -> tuple[str, list[str]]:
    """Extrae texto plano de un CV. `filename` decide el parser por extension."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith((".txt", ".md")):
        try:
            return data.decode("utf-8", "replace").strip(), []
        except Exception as exc:  # noqa: BLE001
            return "", [f"txt fallo: {exc}"]
    # Intento best-effort: probar PDF y luego texto plano.
    text, warns = _extract_pdf(data)
    if text:
        return text, warns
    return data.decode("utf-8", "replace").strip(), [
        f"Extension no reconocida ({filename}); tratado como texto plano."
    ]
